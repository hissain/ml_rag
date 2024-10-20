import time
import faiss
import json
import requests
import numpy as np
import PyPDF2
import docx
import pandas as pd
from tqdm.notebook import tqdm
#from ollama import Ollama

ollama_url_inf = "http://localhost:11434/api/show"
ollama_url_emb = "http://localhost:11434/api/embeddings"
ollama_url_gen = "http://localhost:11434/api/generate"
ollama_model_name = "llama3.2:latest"

VERBOSE = False


# File as URI

def read_pdf(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text()
    return text

def read_docx(file_path):
    doc = docx.Document(file_path)
    text = "\n".join([p.text for p in doc.paragraphs])
    return text

def read_excel(file_path):
    df = pd.read_excel(file_path)
    return df.to_string()

def read_text(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def read_file(file_path):
    if file_path.endswith('.pdf'):
        return read_pdf(file_path)
    elif file_path.endswith('.docx'):
        return read_docx(file_path)
    elif file_path.endswith('.xlsx'):
        return read_excel(file_path)
    elif file_path.endswith('.txt'):
        return read_text(file_path)
    else:
        raise ValueError("Unsupported file format")
    

# File as objects

def read_pdf_obj(file_obj):
    text = ""
    reader = PyPDF2.PdfReader(file_obj)  # File-like object
    for page in reader.pages:
        text += page.extract_text()
    return text

def read_docx_obj(file_obj):
    doc = docx.Document(file_obj)  # File-like object
    text = "\n".join([p.text for p in doc.paragraphs])
    return text

def read_excel_obj(file_obj):
    df = pd.read_excel(file_obj)  # File-like object
    return df.to_string()

def read_text_obj(file_obj):
    return file_obj.read().decode('utf-8')

def read_file_as_obj(file_obj):
    file_name = file_obj.name  # Get the name of the uploaded file
    if file_name.endswith('.pdf'):
        return read_pdf_obj(file_obj)  # Pass file-like object
    elif file_name.endswith('.docx'):
        return read_docx_obj(file_obj)  # Pass file-like object
    elif file_name.endswith('.xlsx'):
        return read_excel_obj(file_obj)  # Pass file-like object
    elif file_name.endswith('.txt'):
        return read_text_obj(file_obj)  # Pass file-like object
    else:
        raise ValueError("Unsupported file format")    

def partition_text(text, max_length):
    sentences = text.split('. ')
    partitions = []
    current_part = []
    current_length = 0
    
    for sentence in tqdm(sentences, desc="Partitioning text"):
        current_length += len(sentence.split())
        current_part.append(sentence)
        
        if current_length > max_length:
            partitions.append('. '.join(current_part))
            current_part = []
            current_length = 0

    if current_part:
        partitions.append('. '.join(current_part))

    return partitions

def get_embedding_shape():
    payload = { "model": ollama_model_name }
    headers = {"Content-Type": "application/json"}
    response = requests.post(ollama_url_inf, headers=headers, data=json.dumps(payload))

    if response.status_code == 200:
        result = response.json()
        if 'model_info' in result and 'llama.embedding_length' in result['model_info']:
            embedding_length = result['model_info']["llama.embedding_length"]
            return embedding_length
        else:
            return 0
    else:
        return 0

def get_embedding(text):
    payload = { "model": ollama_model_name, "prompt": text}
    headers = {"Content-Type": "application/json"}
    response = requests.post(ollama_url_emb, headers=headers, data=json.dumps(payload))

    if response.status_code == 200:
        result = response.json()
        embedding = np.array(result['embedding'])
        return embedding
    else:
        return np.zeros(768)  # (adjust dimension based on model)

def store_in_faiss(partitions):
    dimension = get_embedding_shape()
    index = faiss.IndexFlatL2(dimension)
    doc_vectors = []
    doc_ids = []
    
    for i, partition in tqdm(enumerate(partitions), total=len(partitions), desc="Embedding partitions"):
        embedding = get_embedding(partition)
        index.add(np.array([embedding]))
        doc_vectors.append(embedding)
        doc_ids.append(i)
    
    return index, doc_ids

def retrieve_with_rag(query, faiss_index, doc_ids, partitions, k=2):
    query_embedding = get_embedding(query)
    distances, indices = faiss_index.search(np.array([query_embedding]), k=k)
    retrieved_docs = []
    for i in tqdm(indices[0], desc="Retrieving documents"):
        if i >= len(partitions):
            continue
        doc_id = doc_ids[i]
        retrieved_docs.append(partitions[doc_id])
    combined_docs = "\n".join(retrieved_docs)
    rag_prompt = f"Context:\n{combined_docs}\n\nQuery: {query}\nAnswer:"
    payload = {"model": ollama_model_name, "prompt": rag_prompt, "stream": False}
    response = requests.post(ollama_url_gen, headers={"Content-Type": "application/json"}, 
                             data=json.dumps(payload))
    return response.json()

def ask(query, faiss_index, doc_ids, partitions):
    rag_response = retrieve_with_rag(query, faiss_index, doc_ids, partitions)
    return rag_response["response"]


def main():
    file_path = "./../data/designpattern.pdf"
    text_data = read_file(file_path)
    partitions = partition_text(text_data, max_length=512)
    faiss_index, doc_ids = store_in_faiss(partitions)

if __name__ == '__main__':
    main()